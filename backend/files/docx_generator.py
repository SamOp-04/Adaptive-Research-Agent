from __future__ import annotations

from pathlib import Path

from backend.files.markdown_rendering import (
    MarkdownBlock,
    iter_inline_segments,
    parse_markdown_blocks,
    plain_text_from_markdown,
)
from backend.graph.state import SourceFinding


def write_docx(
    content: str,
    path: str | Path,
    *,
    title: str = "Adaptive Research Report",
    key_findings: list[str] | None = None,
    sources: list[SourceFinding] | None = None,
) -> Path:
    from docx import Document
    from docx.shared import Pt, RGBColor

    destination = Path(path)
    document = Document()
    document.styles["Normal"].font.size = Pt(11)
    document.styles["Normal"].font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    title_heading = document.add_heading(title, level=1)
    for run in title_heading.runs:
        run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

    document.add_heading("Summary", level=2)
    _add_markdown_blocks(document, content)

    document.add_heading("Key Findings", level=2)
    if key_findings:
        for finding in key_findings:
            document.add_paragraph(finding, style="List Bullet")
    else:
        document.add_paragraph("None", style="List Bullet")

    document.add_heading("Sources", level=2)
    if sources:
        for source in sources[:12]:
            _add_source_paragraph(document, source)
    else:
        document.add_paragraph("None", style="List Bullet")

    document.save(destination)
    return destination


def _add_markdown_blocks(document: object, content: str) -> None:
    from docx.shared import Pt

    blocks = parse_markdown_blocks(content)
    if not blocks:
        document.add_paragraph("No content generated.")
        return

    for block in blocks:
        if block.type == "heading":
            heading = document.add_heading("", level=min(max(block.level, 2), 4))
            _add_inline_runs(heading, block.text)
        elif block.type == "paragraph":
            paragraph = document.add_paragraph()
            _add_inline_runs(paragraph, block.text)
        elif block.type == "bullet_list":
            for item in block.items or []:
                paragraph = document.add_paragraph(style="List Bullet")
                _add_inline_runs(paragraph, item)
        elif block.type == "ordered_list":
            for item in block.items or []:
                paragraph = document.add_paragraph(style="List Number")
                _add_inline_runs(paragraph, item)
        elif block.type == "table":
            _add_table(document, block)
        elif block.type == "code":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(block.text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)


def _add_table(document: object, block: MarkdownBlock) -> None:
    headers = block.headers or []
    if not headers:
        return

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = plain_text_from_markdown(header)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row in block.rows or []:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = plain_text_from_markdown(value)


def _add_inline_runs(paragraph: object, text: str) -> None:
    from docx.shared import Pt

    for segment in iter_inline_segments(text):
        run = paragraph.add_run(segment.text)
        run.bold = segment.bold
        run.italic = segment.italic or bool(segment.url)
        if segment.code:
            run.font.name = "Courier New"
        if segment.url:
            url_run = paragraph.add_run(f" ({segment.url})")
            url_run.italic = True
            url_run.font.size = Pt(9)


def _credibility_tier(source: SourceFinding) -> str:
    score = float(source.get("credibility", {}).get("score", 0) or 0)
    if score >= 0.75:
        return "High"
    if score >= 0.5:
        return "Medium"
    return "Low"


def _add_source_paragraph(document: object, source: SourceFinding) -> None:
    from docx.shared import Pt

    domain = source.get("credibility", {}).get("domain") or source.get("url", "")
    tier = _credibility_tier(source)
    paragraph = document.add_paragraph(style="List Bullet")
    tag_run = paragraph.add_run(f"[{tier}] ")
    tag_run.font.size = Pt(10)
    title_run = paragraph.add_run(plain_text_from_markdown(source.get("title", "Untitled source")))
    title_run.bold = True
    title_run.font.size = Pt(10)
    domain_run = paragraph.add_run(f" - {domain}")
    domain_run.font.size = Pt(10)
    url_run = paragraph.add_run(f"\n{source.get('url', '')}")
    url_run.italic = True
    url_run.font.size = Pt(10)
