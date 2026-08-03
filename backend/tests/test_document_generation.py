from __future__ import annotations

from docx import Document

from backend.files.docx_generator import write_docx
from backend.files.markdown_rendering import markdown_to_html, parse_markdown_blocks


def test_markdown_to_html_renders_tables_and_emphasis():
    html = markdown_to_html(
        "### Findings\n\n"
        "| Driver | Impact |\n"
        "| --- | --- |\n"
        "| **Policy** | High |\n\n"
        "---\n\n"
        "Use `plain text`."
    )

    assert "<h3>" in html
    assert "<table>" in html
    assert "<strong>Policy</strong>" in html
    assert "<hr>" in html
    assert "<code>plain text</code>" in html


def test_parse_markdown_blocks_detects_gfm_table():
    blocks = parse_markdown_blocks("| A | B |\n| --- | --- |\n| 1 | 2 |")

    assert len(blocks) == 1
    assert blocks[0].type == "table"
    assert blocks[0].headers == ["A", "B"]
    assert blocks[0].rows == [["1", "2"]]


def test_write_docx_turns_markdown_into_docx_structure(tmp_path):
    path = tmp_path / "report.docx"
    write_docx(
        "### Findings\n\n"
        "This has **bold** text.\n\n"
        "| Driver | Impact |\n"
        "| --- | --- |\n"
        "| Policy | High |\n\n"
        "---",
        path,
        title="Markdown report",
    )

    document = Document(path)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert len(document.tables) == 1
    assert document.tables[0].cell(0, 0).text == "Driver"
    assert document.tables[0].cell(1, 1).text == "High"
    assert "**bold**" not in paragraph_text
    assert "### Findings" not in paragraph_text
    assert "---" not in paragraph_text
    assert "bold" in paragraph_text
